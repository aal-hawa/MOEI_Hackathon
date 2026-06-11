from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("/Users/hassaneltahir/Downloads/Sultan_Rescheduling_Test_Documents")


INK = RGBColor(10, 22, 40)
GOLD = RGBColor(201, 168, 76)
RED = RGBColor(154, 35, 35)
MUTED = RGBColor(91, 104, 124)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
      run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_test_banner(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HACKATHON TEST ONLY - NOT A VALID FINANCIAL, BANK, GOVERNMENT, OR EMPLOYMENT DOCUMENT")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RED

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Synthetic demo data for testing the MOEI arrears rescheduling AI agent only.")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_test_banner(doc)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = INK

    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(21)
    r.font.color.rgb = INK

    s = doc.add_paragraph()
    s.paragraph_format.space_after = Pt(12)
    sr = s.add_run(subtitle)
    sr.font.name = "Calibri"
    sr.font.size = Pt(10)
    sr.font.color.rgb = MUTED

    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_after = Pt(12)
    wr = warning.add_run("MOCK / SAMPLE / FOR UPLOAD TESTING ONLY")
    wr.bold = True
    wr.font.name = "Calibri"
    wr.font.size = Pt(13)
    wr.font.color.rgb = RED
    return doc


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD


def add_key_table(doc: Document, rows: list[tuple[str, str]], widths=(2.35, 4.65)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        set_cell_shading(cells[0], "F2F4F7")
        set_cell_text(cells[0], label, bold=True, color=INK)
        set_cell_text(cells[1], value, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def build_salary_certificate() -> Path:
    doc = setup_doc(
        "Mock Salary Certificate",
        "Fictional employment evidence for Sultan Al Nuaimi's arrears rescheduling test case",
    )
    add_section_heading(doc, "Employee And Employer Details")
    add_key_table(doc, [
        ("Certificate Reference", "MOCK-SAL-2026-0007"),
        ("Issue Date", "08 June 2026"),
        ("Employee Name", "Sultan Al Nuaimi"),
        ("Emirates ID", "784-1992-9876543-1"),
        ("Employer", "Gulf Infrastructure Services LLC - Fictional Demo Employer"),
        ("Position", "Senior Facilities Coordinator"),
        ("Employment Status", "Active, full-time"),
        ("Date of Joining", "15 March 2018"),
    ])

    add_section_heading(doc, "Monthly Income Components")
    add_key_table(doc, [
        ("Basic Salary", "AED 15,200"),
        ("Housing Allowance", "AED 2,800"),
        ("Transport Allowance", "AED 900"),
        ("Gross Monthly Income", "AED 18,900"),
        ("Estimated Net Salary Credit", "AED 17,250"),
    ])

    add_section_heading(doc, "Purpose")
    doc.add_paragraph(
        "This mock certificate is provided only to test the MOEI housing arrears rescheduling AI agent. "
        "It should be uploaded as the salary_certificate document type during the demo."
    )
    add_note(doc, "This document intentionally has no official logo, stamp, signature, or real employer authority.")

    path = OUT_DIR / "01_MOCK_Sultan_Al_Nuaimi_Salary_Certificate.docx"
    doc.save(path)
    return path


def build_income_statement() -> Path:
    doc = setup_doc(
        "Mock Detailed Income And Bank Statement",
        "Fictional six-month income evidence for Sultan Al Nuaimi's rescheduling request",
    )
    add_section_heading(doc, "Account Summary")
    add_key_table(doc, [
        ("Statement Reference", "MOCK-BANK-2026-0043"),
        ("Account Holder", "Sultan Al Nuaimi"),
        ("Account Ending", "4431"),
        ("Statement Period", "01 January 2026 to 31 May 2026"),
        ("Average Salary Credit", "AED 17,250"),
        ("Average Monthly Balance", "AED 8,410"),
    ])

    add_section_heading(doc, "Salary Credits")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(["Month", "Credit Date", "Description", "Amount"]):
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK)
    credits = [
        ("Jan 2026", "28 Jan 2026", "Salary credit - fictional employer", "AED 17,250"),
        ("Feb 2026", "27 Feb 2026", "Salary credit - fictional employer", "AED 17,250"),
        ("Mar 2026", "28 Mar 2026", "Salary credit - fictional employer", "AED 17,250"),
        ("Apr 2026", "28 Apr 2026", "Salary credit - fictional employer", "AED 17,250"),
        ("May 2026", "28 May 2026", "Salary credit - fictional employer", "AED 17,250"),
    ]
    for row in credits:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, color=INK)

    add_section_heading(doc, "Recurring Monthly Obligations")
    add_key_table(doc, [
        ("Existing Housing Loan Installment", "AED 3,000"),
        ("Vehicle Finance", "AED 1,100"),
        ("School Fees Reserve", "AED 950"),
        ("Utilities And Household Commitments", "AED 600"),
        ("Estimated Monthly Obligations", "AED 5,650"),
    ])
    doc.add_paragraph(
        "This mock statement supports stable income verification and should be uploaded as bank_statement "
        "or income_statement during the demo."
    )
    add_note(doc, "This is not a real bank statement and cannot be used for any financial transaction.")

    path = OUT_DIR / "02_MOCK_Sultan_Al_Nuaimi_Income_Bank_Statement.docx"
    doc.save(path)
    return path


def build_hardship_letter() -> Path:
    doc = setup_doc(
        "Mock Arrears Rescheduling Request Letter",
        "Fictional beneficiary explanation for the housing arrears rescheduling demo",
    )
    add_section_heading(doc, "Applicant Request")
    doc.add_paragraph(
        "I, Sultan Al Nuaimi, request rescheduling of my housing loan arrears because temporary family "
        "obligations caused missed installments. My income is stable and I want to continue repayment under "
        "a plan that remains within the MOEI 20% deduction cap."
    )

    add_section_heading(doc, "Requested Plan")
    add_key_table(doc, [
        ("Requested Request Type", "Reschedule arrears"),
        ("Arrears Amount", "AED 42,000"),
        ("Remaining Loan Balance", "AED 360,000"),
        ("Requested Duration", "120 months"),
        ("Target Monthly Installment", "Approximately AED 3,350"),
        ("Monthly Income", "AED 18,900"),
        ("Target Deduction Rate", "17.7%"),
    ])

    add_section_heading(doc, "Applicant Declaration")
    doc.add_paragraph(
        "For this mock demo, the applicant confirms willingness to continue repayment and provide any "
        "additional income evidence requested by the officer or AI assessment workflow."
    )
    add_note(doc, "This request letter is synthetic and should only be used in the local hackathon test app.")

    path = OUT_DIR / "03_MOCK_Sultan_Al_Nuaimi_Rescheduling_Request_Letter.docx"
    doc.save(path)
    return path


def build_cheat_sheet() -> Path:
    doc = setup_doc(
        "Sultan Test Input Cheat Sheet",
        "Use these values in the local app to test the Decision Twin auto-approval path",
    )
    add_section_heading(doc, "Citizen And Applicant Values")
    add_key_table(doc, [
        ("Applicant Name", "Sultan Al Nuaimi"),
        ("Emirates ID", "784-1992-9876543-1"),
        ("Monthly Income", "AED 18,900"),
        ("Employer Type", "Government or semi-government"),
        ("Family Size", "5"),
        ("Income Stability", "Stable"),
    ])

    add_section_heading(doc, "Loan And Arrears Values")
    add_key_table(doc, [
        ("Original Loan Amount", "AED 650,000"),
        ("Remaining Balance", "AED 360,000"),
        ("Current Monthly Installment", "AED 3,000"),
        ("Loan Duration", "240 months"),
        ("Elapsed Months", "96 months"),
        ("Arrears Amount", "AED 42,000"),
        ("Missed Months", "14"),
        ("Delay Days", "120"),
        ("Requested Rescheduling Duration", "120 months"),
        ("Reason Category", "reschedule_arrears"),
    ])

    add_section_heading(doc, "Expected AI Result")
    add_key_table(doc, [
        ("Recommended Amount", "AED 402,000"),
        ("Estimated Installment", "AED 3,350"),
        ("Expected Deduction Rate", "17.7%"),
        ("20% Rule", "Pass"),
        ("Document Completeness", "Pass when salary certificate and bank/income statement are uploaded"),
        ("Decision Twin Path", "Auto path / Route to green clear"),
    ])
    add_note(doc, "Upload document 01 as salary_certificate and document 02 as bank_statement or income_statement.")

    path = OUT_DIR / "00_README_Sultan_Test_Input_Cheat_Sheet.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = [
        build_cheat_sheet(),
        build_salary_certificate(),
        build_income_statement(),
        build_hardship_letter(),
    ]
    print("\n".join(str(path) for path in paths))


if __name__ == "__main__":
    main()
